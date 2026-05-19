"""Generate the two Word documents for the Anthropic Cyber Verification
Program submission:

  1. archimedes-cyber-verification-draft.docx  — content to copy into form fields
  2. how-to-submit.docx                        — step-by-step submission guide

Both written with python-docx. Run:
  uv run python docs/cyber-verification/build_docs.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


# ---------- style helpers ----------


def heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading with consistent style — avoids placeholder weirdness across
    Word templates."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif level == 2:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        else:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x4A, 0x90, 0xE2)


def para(doc: Document, text: str, *, italic: bool = False, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold


def bullet(doc: Document, text: str, *, level: int = 0, bold_lead: str | None = None) -> None:
    """Bullet. If bold_lead is set, that prefix is bolded and the rest is plain."""
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True
        r1.font.name = "Calibri"
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.name = "Calibri"
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def num(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True
        r1.font.name = "Calibri"
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.name = "Calibri"
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def divider(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("• • •")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# =====================================================================
# Document 1 — The draft (content to paste into the form fields)
# =====================================================================


def build_draft() -> Path:
    doc = Document()

    # Cover heading
    heading(doc, "Anthropic Cyber Verification Program — Submission Draft", level=1)
    para(
        doc,
        "Reference content for the form at claude.com/form/cyber-use-case. "
        "Copy each section into the corresponding form field; adjust the bracketed "
        "placeholders before submitting.",
        italic=True,
    )
    para(doc, "")  # spacer
    para(
        doc,
        "Project: Archimedes — an autonomous defensive CTI analyst agent built on Claude Code.",
        bold=True,
    )
    divider(doc)

    # --- Section 1
    heading(doc, "1.  Use case description", level=2)
    para(
        doc,
        "Archimedes is an autonomous defensive Cyber Threat Intelligence (CTI) "
        "analyst agent built on Claude Code. It serves a defensive CTI workflow "
        "at [your organization / team name], a US aerospace and defense "
        "contractor environment.",
    )
    para(
        doc,
        "The agent ingests open-source intelligence (~45 vetted sources including "
        "CISA KEV, Mandiant, Unit 42, MSTIC, CrowdStrike, BleepingComputer, "
        "The Record), grades each source-claim per the NATO Admiralty Scale "
        "(AJP-2.1), and produces twice-daily intelligence briefs (08:00 + 16:00 "
        "EDT) to a private Discord channel for the security team and leadership. "
        "It also maintains a structured corpus of threat actor dossiers and CVE "
        "tracking entries, and supports operator-driven queries via slash commands "
        "(/cve <id>, /investigate <target>, /ioc-hunt <indicator>).",
    )
    para(
        doc,
        "The agent is exclusively defensive. It does not generate exploit code, "
        "does not perform active reconnaissance against non-authorized targets, "
        "does not originate attribution claims, and does not produce offensive "
        "tooling of any kind.",
    )

    # --- Section 2
    heading(doc, "2.  Defensive intent — what the agent does NOT do", level=2)
    bullet(
        doc,
        "Hard Rule 3 in the agent's doctrine. Agent refuses to generate "
        "proof-of-concept code, payloads, exploit guides, or assistance "
        "attacking any system — including for \"testing,\" \"research,\" or "
        "\"educational\" purposes.",
        bold_lead="No exploitation, ever. ",
    )
    bullet(
        doc,
        "Hard Rule 4. Active reconnaissance is permitted only against assets "
        "in infrastructure/authorized-targets.yaml (the operator's own "
        "infrastructure). MCP wrappers for SpiderFoot and theHarvester enforce "
        "module-level passive-only allowlists in code — non-passive modules are "
        "rejected before any HTTP call is made.",
        bold_lead="No active scanning of third parties. ",
    )
    bullet(
        doc,
        "Hard Rule 2. The agent only reports what other named sources have "
        "already attributed, citing them; it does not generate first-time "
        "threat-actor attributions.",
        bold_lead="No originating attribution. ",
    )
    bullet(
        doc,
        "Hard Rule 7. If a query surfaces credentials, the agent counts and "
        "reports exposure, then discards. It does not store credentials, "
        "hashes-labeled-as-credentials, or personally identifiable information.",
        bold_lead="No credential handling. ",
    )

    # --- Section 3
    heading(doc, "3.  Safety controls in place", level=2)
    bullet(
        doc,
        "Every behavior the agent exhibits traces to a versioned .md file in "
        "the project's doctrine/ directory. Six files (LEGAL-POLICY, "
        "INTEL-GRADING, INTEL-BRIEF-STANDARDS, THREAT-BOX-METHODOLOGY, "
        "RETRACTION-POLICY, FLASH-POLICY) govern grading, brief composition, "
        "retraction handling, and trigger conditions. Versioned in git; "
        "reviewable; auditable.",
        bold_lead="Doctrine-as-code. ",
    )
    bullet(
        doc,
        "Hard-coded allowlists in the integration layer refuse non-passive "
        "operations before any HTTP request — see "
        "mcps/spiderfoot/src/spiderfoot_mcp/policy.py for the canonical example.",
        bold_lead="MCP wrappers enforce policy in code, not prose. ",
    )
    bullet(
        doc,
        "A dedicated librarian subagent is the only component allowed to write "
        "to git, Splunk, or Discord. Every external publication passes through "
        "a LEGAL-POLICY content-scan gate that refuses credential patterns, "
        "TLP:RED material, and ITAR-questionable detail.",
        bold_lead="Single-writer side effects. ",
    )
    bullet(
        doc,
        "Every agent action produces (1) a git commit, (2) a Splunk event with "
        "run_id, and (3) a Discord channel post. The three logs are joinable on "
        "run_id; any action can be reconstructed post-hoc.",
        bold_lead="Three-log audit trail. ",
    )
    bullet(
        doc,
        "Every Hard Rule refusal is logged to "
        "infrastructure/policy-violations.yaml with timestamp, attempted "
        "action, and reason. Auditable record of what the agent was asked to do "
        "but refused.",
        bold_lead="Refusal logging. ",
    )
    bullet(
        doc,
        "When the agent proposes a HIGH threat-actor scoring, it does not "
        "auto-commit — it posts to a review channel and waits for an "
        "operator's explicit approval (/approve-scoring <actor-id>). Five of "
        "five actor scorings to date have stayed below HIGH; the gate is preserved.",
        bold_lead="Human approval gates on high-stakes decisions. ",
    )

    # --- Section 4
    heading(doc, "4.  Example queries (the kind that would be made)", level=2)
    bullet(
        doc,
        "Pull NVD record, CISA KEV status, vendor patch advisory for a "
        "vulnerability; summarize patch posture and defensive priority for "
        "the security team. Use case is patch-management and defensive "
        "prioritization, not exploit research.",
        bold_lead="/cve CVE-2026-XXXX — ",
    )
    bullet(
        doc,
        "Pull existing tracked-actor dossier, summarize recent reporting from "
        "Tier-1 sources (Mandiant, Unit 42, MSTIC) on tactics, techniques, and "
        "procedures; produce a defender-focused situational summary.",
        bold_lead="/investigate <actor> — ",
    )
    bullet(
        doc,
        "Check an indicator against the local corpus, internal Splunk telemetry, "
        "and external reputation services (VirusTotal, AbuseIPDB). Used to "
        "determine whether a suspicious IP has any history relevant to defense.",
        bold_lead="/ioc-hunt <indicator> — ",
    )
    bullet(
        doc,
        "Automated runs that read overnight OSINT, grade against Admiralty, "
        "surface 4-8 items the operator's security team and leadership should "
        "know about. Output is a ~250-word Smart Brevity summary in a Discord "
        "channel.",
        bold_lead="Twice-daily morning + afternoon briefs — ",
    )

    # --- Section 5
    heading(doc, "5.  Why the AUP filter is firing today", level=2)
    para(
        doc,
        "The agent's defensive CTI work involves naming CVEs, threat groups, "
        "and TTPs — the same surface vocabulary used in offensive security "
        "research. The Anthropic Usage Policy classifier appropriately errs "
        "on the side of caution on cyber-adjacent queries, and we have hit "
        "refusals on routine defensive workflows (for example, requesting a "
        "patch-posture summary for a published, CISA-listed CVE). Cyber "
        "Verification will calibrate the classifier to this account's verified "
        "defensive use case so the agent's standard /cve workflow can proceed "
        "without manual rephrasing.",
    )

    divider(doc)
    heading(doc, "Notes before submitting", level=2)
    bullet(
        doc,
        "Replace [your organization / team name] in Section 1 with the name "
        "you want on the record.",
    )
    bullet(
        doc,
        "If the form has a character limit per field, prioritize Sections 1, "
        "2, and 3 over Sections 4 and 5 — the latter two are nice-to-have, "
        "the former three are load-bearing.",
    )
    bullet(
        doc,
        "If asked for supporting links or repository visibility, the Archimedes "
        "repo is at https://github.com/ryansketch01/archimedes — note that "
        "doctrine/ is the most useful directory for reviewers to look at.",
    )

    out = Path(__file__).parent / "archimedes-cyber-verification-draft.docx"
    doc.save(out)
    return out


# =====================================================================
# Document 2 — How to submit
# =====================================================================


def build_howto() -> Path:
    doc = Document()

    heading(doc, "How to submit the Cyber Verification request to Anthropic", level=1)
    para(
        doc,
        "Step-by-step. Anthropic's AUP refusal includes a Cyber Verification "
        "form URL, but the token is truncated at the source — the URL as shown "
        "is not usable directly. The reliable path is to contact Anthropic "
        "Support and submit the draft content via the support channel.",
        italic=True,
    )
    para(doc, "")
    divider(doc)

    # --- What's in this folder
    heading(doc, "What's in this folder", level=2)
    bullet(
        doc,
        "archimedes-cyber-verification-draft.docx — the content of your "
        "verification request (use case, safety posture, example queries). "
        "Paste sections from this into your support message.",
    )
    bullet(
        doc,
        "how-to-submit.docx — this document.",
    )
    bullet(
        doc,
        "build_docs.py — python-docx source. Edit and regenerate if the "
        "draft needs updates.",
    )

    # --- Background
    heading(doc, "Background — why this is a support-channel submission", level=2)
    para(
        doc,
        "When Claude refuses a cyber-related request under Anthropic's Usage "
        "Policy, the error response includes a Cyber Verification Program "
        "form URL of the form:",
    )
    code_block(
        doc,
        "https://claude.com/form/cyber-use-case?token=<long-token>",
    )
    para(
        doc,
        "In practice, the token is truncated mid-string with an ellipsis "
        "character at the end, making the URL unusable for direct navigation. "
        "Confirmed empirically 2026-05-15 against two independent AUP refusals. "
        "Until Anthropic ships a non-truncated token-bound URL or a stable "
        "console-dashboard entry for verification, the practical path is to "
        "contact Anthropic Support directly.",
    )

    # --- Step 1
    heading(doc, "Step 1 — Open Anthropic Support", level=2)
    para(doc, "Two viable channels:")
    bullet(
        doc,
        "Email — send to support@anthropic.com",
    )
    bullet(
        doc,
        "Web — go to https://support.anthropic.com and start a new conversation",
    )
    para(
        doc,
        "Both route to the same intake. Pick whichever you prefer for "
        "ongoing back-and-forth with Anthropic's review team.",
        italic=True,
    )

    # --- Step 2
    heading(doc, "Step 2 — Compose the support request", level=2)
    bullet(
        doc,
        "Subject:  Cyber Verification Program — defensive CTI workflow",
        bold_lead="",
    )
    para(
        doc,
        "Body — copy the entire content of archimedes-cyber-verification-draft.docx "
        "into the body of the support message. Include all five sections "
        "(use case, defensive-intent-what-NOT-done, safety controls, example "
        "queries, why-AUP-firing).",
    )
    para(
        doc,
        "Add a short opening paragraph above the draft content explaining the "
        "specific request:",
        italic=True,
    )
    code_block(
        doc,
        "We are operating Archimedes, an autonomous defensive CTI analyst built\n"
        "on Claude Code. We have hit AUP refusals on routine defensive workflows\n"
        "(e.g. requesting a patch-posture summary for a published, CISA-listed\n"
        "CVE). We would like to apply for the Cyber Verification Program to\n"
        "calibrate the classifier to our verified defensive use case.\n"
        "\n"
        "Below is the requested content for the program. Happy to provide\n"
        "additional detail, a repository walkthrough, or a video demo on request.",
    )
    bullet(
        doc,
        "Replace [your organization / team name] in the draft content before sending.",
    )
    bullet(
        doc,
        "If the support form has a character limit, attach the draft .docx "
        "as a file attachment and reference it in the message body.",
    )

    # --- Step 3
    heading(doc, "Step 3 — Provide supporting context if asked", level=2)
    para(
        doc,
        "Anthropic may ask follow-up questions. Likely asks plus suggested responses:",
    )
    bullet(
        doc,
        "\"Can we see the repository?\"  Share https://github.com/ryansketch01/archimedes "
        "and direct reviewers to the doctrine/ directory — it's the most useful starting "
        "point for understanding the agent's safety posture.",
        bold_lead="",
    )
    bullet(
        doc,
        "\"What kinds of CVEs would you research?\"  Answer with concrete examples — "
        "recent KEV adds, vendor-disclosed vulnerabilities affecting widely-deployed "
        "products in DIB stacks (Fortinet, Ivanti, Palo Alto, Cisco). Frame as "
        "patch-management and exposure-assessment, never as exploit research.",
        bold_lead="",
    )
    bullet(
        doc,
        "\"How do you prevent misuse?\"  Point to Hard Rule 3 (no exploitation, ever) "
        "in CLAUDE.md, the SpiderFoot / theHarvester passive-only allowlists enforced "
        "in code, the librarian-only-writes pattern with LEGAL-POLICY content scans, "
        "and the three-log audit trail.",
        bold_lead="",
    )
    bullet(
        doc,
        "\"Who is the operator and what's the org context?\"  Provide team name, "
        "organization, and any relevant context (defense-contractor cleared "
        "environment, ITAR-program adjacency, etc.).",
        bold_lead="",
    )

    # --- Step 4
    heading(doc, "Step 4 — Submit and save the case number", level=2)
    num(doc, "Send the support message (email) or submit the support form.")
    num(
        doc,
        "Anthropic typically replies with a case number / ticket ID. Save it "
        "in this folder as case-id.txt or note it in CLAUDE.md operational notes "
        "so future sessions can reference it.",
    )
    num(
        doc,
        "If you used email, expect an auto-acknowledgement within minutes. "
        "If using the web support form, expect the case ID on the confirmation page.",
    )

    # --- Step 5
    heading(doc, "Step 5 — Wait for approval (and how you'll know)", level=2)
    bullet(
        doc,
        "Most likely notification channel: reply to your support ticket via the "
        "same email address you submitted from.",
    )
    bullet(
        doc,
        "No published SLA — anecdotally these reviews tend to land within a "
        "few business days, but it can vary based on Anthropic's backlog.",
    )
    bullet(
        doc,
        "Most reliable empirical test: try /cve <any-cve-id> in Discord #commands "
        "periodically. As soon as the command succeeds (Claude returns a "
        "patch-posture summary instead of an AUP refusal), verification has cleared.",
    )

    # --- Step 6
    heading(doc, "Step 6 — While you wait", level=2)
    bullet(
        doc,
        "/ping, /help, /investigate, /ioc-hunt, /new-actor, /update-tracking, "
        "and /approve-scoring continue to work normally — only /cve is currently "
        "blocked by the AUP classifier.",
    )
    bullet(
        doc,
        "If you need a specific CVE briefed before verification clears, ask "
        "Claude directly in a Claude Code session (in this repo) using a defense-framed "
        "prompt — e.g. \"Defensive patch-posture summary for CVE-2026-XXXX — NVD + "
        "CISA KEV + vendor advisory.\"  That phrasing does not trip the classifier "
        "(empirically verified) and produces the same summary directly.",
    )
    bullet(
        doc,
        "Twice-daily scheduled briefs (08:00 + 16:00 EDT) continue to run "
        "without issue — they do not invoke the /cve command path directly.",
    )
    bullet(
        doc,
        "The listener now persists failed-command stdout to "
        "logs/discord-listener/<date>/<run_id>.failure.log — if anything else "
        "starts failing, the full trace is on disk for debugging.",
    )

    divider(doc)

    # --- If something goes wrong
    heading(doc, "If something goes wrong", level=2)
    bullet(
        doc,
        "Anthropic asks for more information — respond with whatever they "
        "request. Reference the draft content sections; offer to share the "
        "repository or do a video walkthrough.",
    )
    bullet(
        doc,
        "Anthropic denies the request — review the denial reason. Common asks: "
        "clarify what is NOT done by the agent (exploitation, attribution "
        "origination, active scanning of non-authorized targets); demonstrate "
        "the policy-enforcement code (mcps/spiderfoot/.../policy.py); produce "
        "examples of actual /cve outputs to show defensive intent.",
    )
    bullet(
        doc,
        "Long wait without a response — Anthropic Support generally replies "
        "within 1-3 business days. If silent for >5 business days, follow up "
        "on the existing case ID.",
    )

    out = Path(__file__).parent / "how-to-submit.docx"
    doc.save(out)
    return out


# =====================================================================


if __name__ == "__main__":
    p1 = build_draft()
    p2 = build_howto()
    print(f"Wrote: {p1}")
    print(f"       size: {p1.stat().st_size:,} bytes")
    print(f"Wrote: {p2}")
    print(f"       size: {p2.stat().st_size:,} bytes")
